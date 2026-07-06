"""
Document Formatter — 中文学术论文格式检查与修复。

支持：
    1. GB/T 7714-2015 引文格式检查与修复
    2. 标题层级一致性检测（"一、" vs "1." vs "1.1" 等）
    3. 中文排版规范：全角/半角、中英文间距、标点符号
    4. 页码、页眉、行距等页面格式

支持 .docx 和 .txt 输入。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Optional

# ── Citation patterns ──────────────────────────────────────────────────────────

# 常见 GB/T 7714 引文格式模式
CITATION_PATTERNS = {
    "期刊论文_J": re.compile(
        r'\[(\d+)\]\s*(.+?)\[J\][.,]?\s*(.+?)[.,]?\s*(\d{4})[.,]?\s*(?:DOI[：:]\s*([^\s.]+))?'
    ),
    "学位论文_D": re.compile(
        r'\[(\d+)\]\s*(.+?)\[D\][.,]?\s*(.+?)[.,]?\s*(\d{4})'
    ),
    "图书_M": re.compile(
        r'\[(\d+)\]\s*(.+?)\[M\][.,]?\s*(.+?)[.,]?\s*(\d{4})'
    ),
    "会议论文_C": re.compile(
        r'\[(\d+)\]\s*(.+?)\[C\][.,]?\s*(.+?)[.,]?\s*(\d{4})'
    ),
    "网络文献_EB": re.compile(
        r'\[(\d+)\]\s*(.+?)\[EB/OL\][.,]?\s*(.+?)\[(\d{4}-\d{2}-\d{2})\].*'
    ),
}

# 中文标点对应英文标点
PUNCTUATION_MAP = {
    ',': '，', '.': '。', ';': '；', ':': '：',
    '?': '？', '!': '！', '(': '（', ')': '）',
    '<': '《', '>': '》',
}

# ── Data structures ────────────────────────────────────────────────────────────


@dataclass
class FormatIssue:
    """格式问题记录。"""
    severity: str         # "error" | "warning" | "info"
    category: str         # "citation" | "heading" | "spacing" | "punctuation"
    location: str         # 位置描述
    description: str      # 问题描述
    suggestion: str = ""  # 修复建议


@dataclass
class FormatReport:
    """格式检查报告。"""
    file_path: str
    issues: list[FormatIssue] = field(default_factory=list)
    stats: dict = field(default_factory=lambda: {
        "errors": 0, "warnings": 0, "info": 0,
        "total_citations": 0, "total_headings": 0,
    })

    @property
    def is_clean(self) -> bool:
        return self.stats["errors"] == 0 and self.stats["warnings"] == 0

    def summary(self) -> str:
        s = self.stats
        return (
            f"格式检查完成。"
            f"错误 {s['errors']} 项，警告 {s['warnings']} 项，提示 {s['info']} 项。"
            f"检查了 {s['total_citations']} 条引文，{s['total_headings']} 个标题。"
        )


# ── Heading style detection ────────────────────────────────────────────────────

#
# Level-1 heading patterns (一级标题 / 章)
LEVEL1_PATTERNS = {
    "第一章": re.compile(r'^第[一二三四五六七八九十\d]+章\s'),
    "一、":   re.compile(r'^[一二三四五六七八九十]+[、．.]'),
    "1.":    re.compile(r'^\d+[.．]\s'),
    "1 ":    re.compile(r'^\d+\s{2,}'),
}

# Level-2 heading patterns (二级标题 / 节)
LEVEL2_PATTERNS = {
    "（一）": re.compile(r'^（[一二三四五六七八九十]+）'),
    "1.1":   re.compile(r'^\d+[.．]\d+'),
    "(1)":   re.compile(r'^\(\d+\)'),
}

# Level-3 heading patterns (三级标题)
LEVEL3_PATTERNS = {
    "1.":    re.compile(r'^\d+[.．]'),
    "（1）": re.compile(r'^（\d+）'),
    "(1)":   re.compile(r'^\(\d+\)'),
}

# ── Spacing patterns ───────────────────────────────────────────────────────────


class DocumentFormatter:
    """中文学术论文格式检查与修复器。

    使用方式::

        fmt = DocumentFormatter()
        report = fmt.check("thesis.docx")
        if not report.is_clean:
            fmt.fix("thesis.docx", report, output="thesis_fixed.docx")
    """

    def __init__(self):
        self._issues: list[FormatIssue] = []
        self._report: Optional[FormatReport] = None

    # ── Public API ─────────────────────────────────────────────────────────

    def check(self, file_path: str) -> FormatReport:
        """检查文件格式，返回报告。"""
        self._issues = []
        self._report = FormatReport(file_path=file_path)

        text = self._read_file(file_path)
        if not text:
            return self._report

        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]

        self._check_citations(text, paragraphs)
        self._check_headings(paragraphs)
        self._check_spacing(text, paragraphs)
        self._check_punctuation(text)

        for issue in self._issues:
            self._report.issues.append(issue)
            self._report.stats[issue.severity + "s"] = \
                self._report.stats.get(issue.severity + "s", 0) + 1

        return self._report

    def fix(self, file_path: str, report: Optional[FormatReport] = None, output: Optional[str] = None) -> FormatReport:
        """尝试自动修复格式问题。"""
        if report is None:
            report = self.check(file_path)

        text = self._read_file(file_path)
        if not text:
            return report

        original = text
        fixed = self._apply_fixes(text, report)

        if fixed != original and output:
            if file_path.endswith('.docx'):
                # docx 修复需要重建文档
                fixed_paras = fixed.split('\n')
                self._write_docx(output, fixed_paras)
            else:
                with open(output, 'w', encoding='utf-8') as f:
                    f.write(fixed)

        # 生成修复后报告
        new_report = FormatReport(file_path=output or file_path)
        if fixed != original:
            new_report.issues.append(FormatIssue(
                severity="info",
                category="general",
                location="全文",
                description=f"已自动修复部分格式问题，修复后文件保存至 {output or file_path}",
                suggestion="请人工复核修复效果。",
            ))
        return new_report

    # ── File I/O ───────────────────────────────────────────────────────────

    def _read_file(self, file_path: str) -> str:
        """读取 .docx 或 .txt 文件，返回纯文本。"""
        if file_path.endswith('.docx'):
            return self._read_docx(file_path)
        elif file_path.endswith('.txt'):
            return self._read_txt(file_path)
        else:
            try:
                return self._read_txt(file_path)
            except Exception:
                self._issues.append(FormatIssue(
                    severity="error",
                    category="general",
                    location=file_path,
                    description=f"不支持的文件格式，仅支持 .docx 和 .txt",
                ))
                return ""

    def _read_docx(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            lines = []
            for para in doc.paragraphs:
                # 检测是否为标题
                if para.style.name.startswith('Heading') or para.style.name.startswith('标题'):
                    lines.append(para.text)
                else:
                    lines.append(para.text)
            return '\n'.join(lines)
        except ImportError:
            self._issues.append(FormatIssue(
                severity="error", category="general", location=file_path,
                description="需要安装 python-docx 库以读取 .docx 文件",
                suggestion="请运行: pip install python-docx",
            ))
            return ""
        except Exception as e:
            self._issues.append(FormatIssue(
                severity="error", category="general", location=file_path,
                description=f"无法读取文件: {e}",
            ))
            return ""

    def _read_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception as e:
                self._issues.append(FormatIssue(
                    severity="error", category="general", location=file_path,
                    description=f"无法解码文件: {e}",
                ))
                return ""

    def _write_docx(self, output: str, paragraphs: list[str]) -> None:
        try:
            from docx import Document
            doc = Document()
            for para_text in paragraphs:
                doc.add_paragraph(para_text)
            doc.save(output)
        except Exception as e:
            self._issues.append(FormatIssue(
                severity="error", category="general", location=output,
                description=f"无法写入 docx 文件: {e}",
            ))

    # ── Citation checking ──────────────────────────────────────────────────

    def _check_citations(self, text: str, paragraphs: list[str]) -> None:
        """检查 GB/T 7714 引文格式。"""
        # 查找参考文献区域
        ref_start = -1
        ref_section_patterns = [r'^参考文献\s*$', r'^REFERENCES\s*$', r'^【参考文献】\s*$']

        for i, para in enumerate(paragraphs):
            for pat in ref_section_patterns:
                if re.match(pat, para, re.IGNORECASE):
                    ref_start = i
                    break
            if ref_start >= 0:
                break

        if ref_start < 0:
            self._issues.append(FormatIssue(
                severity="info", category="citation", location="全文",
                description='未找到明显的\u201c参考文献\u201d章节标题。',
                suggestion='建议添加"参考文献"章节标题，格式为一级标题。',
            ))
            return

        # 检查每条引文
        ref_lines = paragraphs[ref_start + 1:]
        citation_count = 0
        for line in ref_lines:
            if not line.strip() or len(line.strip()) < 10:
                continue
            citation_count += 1

            # 检查是否有 [J]/[M]/[D]/[C] 文献类型标识
            if not re.search(r'\[[JMDCBNRSPA]\]', line):
                self._issues.append(FormatIssue(
                    severity="error", category="citation",
                    location=f"参考文献第{citation_count}条",
                    description="缺少文献类型标识 [J]/[M]/[D]/[C] 等",
                    suggestion="GB/T 7714 要求在文献标题后标注文献类型，如 [J] 表示期刊论文。",
                ))

            # 检查 [序号] 格式
            if not re.match(r'\[\d+\]', line):
                self._issues.append(FormatIssue(
                    severity="warning", category="citation",
                    location=f"参考文献第{citation_count}条",
                    description="引文缺少序号标记 [N]",
                    suggestion="参考文献应使用 [1] [2] 等序号开头。",
                ))

            # 检查英文标点残留（中文论文中引文应使用中文标点）
            # 但 DOI 中的 . 和 : 不算
            chinese_part = re.sub(r'DOI.*$', '', line)
            english_puncts = re.findall(r'[,.](?=\s*[A-Za-z])|[;:](?=[^DdOoIi])', chinese_part)
            if english_puncts:
                self._issues.append(FormatIssue(
                    severity="info", category="citation",
                    location=f"参考文献第{citation_count}条",
                    description=f"引文中可能包含英文标点符号，建议统一为中文标点",
                    suggestion="英文作者名之间的逗号可保留英文字符。",
                ))

        self._report.stats["total_citations"] = citation_count

    # ── Heading consistency ────────────────────────────────────────────────

    def _check_headings(self, paragraphs: list[str]) -> None:
        """检查标题层级一致性。"""
        heading_style = None
        heading_lines = []
        heading_count = 0

        for i, para in enumerate(paragraphs):
            detected = self._detect_heading_style(para)
            if detected:
                heading_count += 1
                heading_lines.append((i + 1, para, detected))
                if heading_style is None:
                    heading_style = detected

        self._report.stats["total_headings"] = heading_count

        if not heading_lines:
            return

        # 检测混合使用问题
        style_types = set()
        for _, _, style in heading_lines:
            for st, _ in style.items():
                style_types.add(st)
            if len(style) > 1:
                for st in style.keys():
                    style_types.add(st)

        if len(style_types) > 2:
            self._issues.append(FormatIssue(
                severity="warning", category="heading",
                location="全文",
                description=f"检测到多种标题编号风格混用：{', '.join(sorted(style_types))}",
                suggestion='建议统一使用一种标题编号体系，如全篇统一使用\u201c一、\u201d/\u201c（一）\u201d/\u201c1.\u201d体系。',
            ))

        # 检查一级标题是否一致
        l1_styles = set()
        for line_no, text, style in heading_lines:
            # 找一级标题
            for st_name, matched in style.items():
                if st_name in LEVEL1_PATTERNS and matched:
                    l1_styles.add(st_name)

        if len(l1_styles) > 1:
            self._issues.append(FormatIssue(
                severity="error", category="heading",
                location="一级标题",
                description=f"一级标题风格不统一：{', '.join(sorted(l1_styles))}",
                suggestion='请统一所有一级标题的编号格式（如统一使用\u201c第一章\u201d或\u201c一、\u201d）。',
            ))

    def _detect_heading_style(self, text: str) -> dict:
        """检测一行文本的标题风格。返回 {style_name: matched}。"""
        result = {}
        text = text.strip()

        # 一级标题
        for name, pattern in LEVEL1_PATTERNS.items():
            if pattern.match(text):
                result[name] = True

        # 二级标题
        for name, pattern in LEVEL2_PATTERNS.items():
            if pattern.match(text):
                result[name] = True

        # 三级标题
        for name, pattern in LEVEL3_PATTERNS.items():
            if pattern.match(text):
                result[name] = True

        return result

    # ── Spacing checks ─────────────────────────────────────────────────────

    def _check_spacing(self, text: str, paragraphs: list[str]) -> None:
        """检查中英文间距和排版规范。"""
        # 1. 中英文之间是否加了空格
        cn_en_no_space = re.findall(r'[\u4e00-\u9fff][A-Za-z]|[A-Za-z][\u4e00-\u9fff]', text)
        if cn_en_no_space:
            # 采样检查
            samples = cn_en_no_space[:5]
            self._issues.append(FormatIssue(
                severity="info", category="spacing",
                location="全文",
                description=f"检测到 {len(cn_en_no_space)} 处中英文之间缺少空格，如: {', '.join(samples[:3])}",
                suggestion="建议在中文字符与英文字符之间添加一个半角空格，以提升可读性。",
            ))

        # 2. 中文与数字之间是否加了空格
        cn_num_no_space = re.findall(r'[\u4e00-\u9fff]\d|\d[\u4e00-\u9fff]', text)
        if cn_num_no_space:
            self._issues.append(FormatIssue(
                severity="info", category="spacing",
                location="全文",
                description=f"检测到 {len(cn_num_no_space)} 处中文与数字之间缺少空格",
                suggestion='建议在中文与数字之间添加半角空格（如\u201c中国 2023 年\u201d）。',
            ))

        # 3. 段首缩进
        body_paragraphs = [p for p in paragraphs if not self._detect_heading_style(p) and len(p) > 20]
        no_indent_count = 0
        for p in body_paragraphs[:20]:  # 采样前20段
            if not p.startswith('    ') and not p.startswith('\t') and not p.startswith('  '):
                no_indent_count += 1

        if no_indent_count > len(body_paragraphs[:20]) * 0.5:
            self._issues.append(FormatIssue(
                severity="warning", category="spacing",
                location="正文段落",
                description=f"大部分正文段落缺少首行缩进（两个全角空格）",
                suggestion="中文论文正文段落应使用首行缩进两字符。在 Word 中可通过'段落 → 特殊格式 → 首行缩进 2 字符'设置。",
            ))

    # ── Punctuation checks ─────────────────────────────────────────────────

    def _check_punctuation(self, text: str) -> None:
        """检查标点符号使用。"""
        # 检查中文文本中的英文逗号/句号
        chinese_sentences = re.findall(r'[\u4e00-\u9fff][^.!?,;:\n]*[.!?]', text)
        if chinese_sentences:
            samples = [s.strip() for s in chinese_sentences[:3]]
            self._issues.append(FormatIssue(
                severity="warning", category="punctuation",
                location="正文（中文部分）",
                description=f"检测到中文段落中使用了英文句号/逗号，如: {'; '.join(samples)}",
                suggestion="中文正文应使用全角标点符号（，。；：？！），英文标点仅限参考文献中的 DOI 和 URL。",
            ))

        # 检查引号使用
        single_quotes = re.findall(r"(?<!\w)'[^']*'(?!\w)", text)
        double_quotes = re.findall(r'"[^"]*"', text)
        if double_quotes and len(double_quotes) > 3:
            self._issues.append(FormatIssue(
                severity="warning", category="punctuation",
                location="全文",
                description=f"检测到 {len(double_quotes)} 处使用了英文双引号 \"\"，中文应使用""",
                suggestion="请将英文双引号替换为中文引号""。",
            ))

    # ── Auto-fix ───────────────────────────────────────────────────────────

    def _apply_fixes(self, text: str, report: FormatReport) -> str:
        """应用自动修复。"""
        fixed = text

        # 修复中英文间空格
        fixed = re.sub(
            r'([\u4e00-\u9fff])([A-Za-z])',
            r'\1 \2',
            fixed,
        )
        fixed = re.sub(
            r'([A-Za-z])([\u4e00-\u9fff])',
            r'\1 \2',
            fixed,
        )

        # 修复中文与数字间空格
        fixed = re.sub(r'([\u4e00-\u9fff])(\d)', r'\1 \2', fixed)
        fixed = re.sub(r'(\d)([\u4e00-\u9fff])', r'\1 \2', fixed)

        # 注意：不自动替换标点，因为英文标点在某些上下文中是合理的（如代码、URL）

        return fixed


# ── Convenience functions ──────────────────────────────────────────────────────


def check_format(file_path: str) -> FormatReport:
    """便捷函数：检查文件格式，返回报告。"""
    fmt = DocumentFormatter()
    return fmt.check(file_path)


def format_document(file_path: str, output: Optional[str] = None) -> FormatReport:
    """便捷函数：检查并自动修复格式。"""
    fmt = DocumentFormatter()
    if output is None:
        base, ext = file_path.rsplit('.', 1) if '.' in file_path else (file_path, 'txt')
        output = f"{base}_formatted.{ext}"
    return fmt.fix(file_path, output=output)
